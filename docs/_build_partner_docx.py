"""Render a partner-friendly DOCX version of the investigation report."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORT_DIR = r"C:\Users\gaomi\Downloads\Telegram Desktop\web-ban-hang-chinh-thuc\web-lms-chinh-thuc\_worktrees\v2-lms-fix\docs"
OUT = REPORT_DIR + r"\LESSON_NAVIGATION_PERFORMANCE_INVESTIGATION_PARTNER.docx"

doc = Document()

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def H(text, level=1, color=(0x2A, 0x4B, 0x2A)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h

def P(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p

def BULLET(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")

def NUM(items):
    for it in items:
        doc.add_paragraph(it, style="List Number")

def CALLOUT(title, body):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    p1 = cell.paragraphs[0]
    r1 = p1.add_run("🔍 " + title)
    r1.bold = True
    r1.font.size = Pt(11)
    cell.add_paragraph(body)

# Title block
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("BÁO CÁO ĐIỀU TRA")
tr.bold = True
tr.font.size = Pt(18)
tr.font.color.rgb = RGBColor(0x2A, 0x4B, 0x2A)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Hiện tượng chuyển bài chậm 5–6 giây trên trang học viên")
sr.font.size = Pt(13)
sr.font.color.rgb = RGBColor(0x4A, 0x3E, 0x3B)
sr.italic = True

P("\nPhiên bản: báo cáo cho đối tác (không kỹ thuật) — dựa trên điều tra ngày 2026-07-19.", italic=True, size=10)
P("Người chuẩn bị: đội kỹ thuật Vercel-LMS. Mục tiêu: giúp người không chuyên kỹ thuật nắm được vấn đề và đề xuất hướng xử lý.\n", italic=True, size=10)

# 1. Tóm tắt
H("1. Hiện tượng đang gặp là gì?", 1)
P("Khi học viên đang xem một bài học và bấm nút “Bài Tiếp Theo” hoặc “Bài Trước”, trang web:")
BULLET([
    "mất khoảng 5–6 giây mới hiển thị nội dung bài mới;",
    "trong thời gian đó có một biểu tượng quay tròn (spinner) che toàn màn hình;",
    "toàn bộ giao diện phải tải lại từ đầu, không phải chỉ phần nội dung đổi.",
])
P("Hiện tượng này lặp lại cho MỌI lần chuyển bài, không riêng một bài nào. Học viên đã phản ánh rằng nó gây khó chịu và đôi khi tưởng trang bị lỗi.\n")

CALLOUT(
    "Hiểu đơn giản",
    "Giống như khi bạn đang xem một trang sách, thay vì chỉ lật sang trang sau, cả quyển sách phải được in lại rồi gửi lại cho bạn. Mỗi lần chuyển bài, hệ thống đang làm việc đó: không tận dụng được những gì đã chuẩn bị sẵn."
)

# 2. Tại sao lại chậm
H("2. Tại sao lại chậm?", 1)
P("Điều tra cho thấy mỗi lần bấm nút chuyển bài, hệ thống đang làm 5 việc nối tiếp nhau. Tổng thời gian của 5 việc đó chính là 5–6 giây mà học viên phải chờ.\n")

H("5 việc nối tiếp gây ra chậm", 2)
NUM([
    "Tải lại toàn bộ trang web (HTML, CSS, hình ảnh mặc định, font chữ) — mất khoảng 0,4–0,5 giây.",
    "Tải một thư viện giao diện tên là Tailwind khoảng 407 KB chỉ để hiển thị trang — mất khoảng 0,4–0,5 giây.",
    "Hỏi máy chủ “bài học này gồm những gì?” — khoảng 0,3–1,3 giây.",
    "Hỏi máy chủ “toàn bộ danh sách bài trong khóa này là gì?” — khoảng 0,3–0,7 giây.",
    "Đợi máy chủ tải nội dung công thức từ Google Drive/Docs cho TẤT CẢ các bài trong khóa — khoảng 1–3 giây, tùy khóa.",
])

P("\nĐáng chú ý: ở bước 5, máy chủ đang tải công thức của tất cả các bài trong cùng khóa học, dù học viên chỉ cần xem một bài duy nhất.\n")

H("Vì sao không phải do mạng hay Supabase chậm?", 2)
P("Để tách bạch rõ nguyên nhân, đội kỹ thuật đã đo từng tầng riêng biệt. Kết quả cho thấy:")
BULLET([
    "Đường truyền mạng (DNS + TLS) chỉ chiếm ~90 ms — đây không phải nút thắt.",
    "Trình duyệt không bị chậm — HTML chỉ ~64 KB, parse rất nhanh.",
    "Supabase không phải lúc nào cũng đụng đến; phần lớn thời gian chờ rơi vào hệ thống làm nhiều việc dư thừa.",
    "Google Drive chỉ là một phần nhỏ trong tổng thời gian, không phải nguyên nhân chính.",
])
P("\nNói cách khác: vấn đề nằm ở THIẾT KẾ luồng xử lý, không phải một dịch vụ cụ thể nào đang chậm.\n")

H("Các thiết kế cũ đang gây lãng phí", 2)
BULLET([
    "Spinner (biểu tượng quay tròn) chỉ tắt được khi TOÀN Bộ danh sách bài + công thức đã sẵn sàng, dù học viên chỉ cần bài hiện tại.",
    "Mỗi lần chuyển bài là một lần tải lại TOÀN BỘ trang, không tận dụng dữ liệu đã có.",
    "Khi tải danh sách bài, máy chủ vô tình tải luôn công thức của TẤT CẢ các bài kèm theo, không riêng bài đang xem.",
    "Trình duyệt bị cấm lưu cache (no-store) vì một cấu hình chung toàn site, nên cứ mỗi lần lại tải lại từ đầu.",
])

# 3. Tác động
H("3. Tác động tới học viên và vận hành", 1)

H("Với học viên", 2)
BULLET([
    "Trải nghiệm khó chịu: 5–6 giây cho mỗi lần chuyển bài, gấp 5–10 lần bình thường.",
    "Tăng tỉ lệ thoát trang, đặc biệt trên di động (mạng chậm hơn).",
    "Trên các khóa có nhiều bài, hiện tượng nặng hơn (do phải tải thêm công thức ở mỗi lần).",
    "Khó tương tác: học viên có thể bấm nhầm khi spinner chưa tắt, dẫn tới nhảy hai bài.",
])

H("Với vận hành", 2)
BULLET([
    "Máy chủ Supabase và Google Drive bị gọi nhiều hơn cần thiết, tốn chi phí API và quota.",
    "Nguy cơ nghẽn khi nhiều học viên chuyển bài cùng lúc (đầu giờ học, cuối tuần).",
    "Khó mở rộng khi thêm bài học hoặc khóa mới — chi phí tăng theo số bài.",
])

# 4. Đề xuất
H("4. Đề xuất xử lý (3 mức độ)", 1)
P("Dưới đây là các đề xuất, xếp theo mức độ tác động và độ phức tạp. Tất cả đều CHƯA được triển khai, chờ đối tác phê duyệt.\n")

H("Mức A — Thay đổi nhỏ, rủi ro thấp (ưu tiên làm trước)", 2)
P("Mục tiêu: cải thiện cảm nhận ngay trong 1–2 tuần, không cần thay đổi kiến trúc.\n")

A = [
    ("A1. Tắt hiệu hình quay sớm hơn",
     "Hiện tại spinner chờ cả danh sách bài tải xong mới tắt. Có thể để nội dung bài hiện tại hiện ra ngay khi tải xong bài đó; danh sách bài sẽ hiện sau. Cảm nhận nhanh hơn rõ rệt mà không tốn thêm chi phí máy chủ."),
    ("A2. Self-host (tự phục vụ) thư viện giao diện",
     "Thay vì mỗi lần chuyển bài lại tải thư viện Tailwind 407 KB từ CDN bên ngoài, đóng gói sẵn vào trang web. Tiết kiệm 0,4 giây mỗi lần."),
    ("A3. Tải sẵn bài kế tiếp/bài trước ngầm",
     "Khi học viên đang xem một bài, hệ thống âm thầm chuẩn bị sẵn dữ liệu cho bài trước/sau. Bấm nút là hiện ra gần như ngay lập tức."),
]
for t, d in A:
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x2A, 0x4B, 0x2A)
    doc.add_paragraph(d, style="Body Text")

H("Mức B — Thay đổi luồng dữ liệu, cần test kỹ hơn", 2)
P("Mục tiêu: giảm thời gian thực sự từ phía máy chủ, không chỉ cảm nhận.\n")

B_items = [
    ("B1. Tách nhỏ danh sách bài và công thức",
     "Hiện nay máy chủ tải công thức của TẤT CẢ bài trong khóa mỗi lần chuyển bài. Đề xuất tách ra: danh sách bài thì nhẹ, công thức chỉ tải khi học viên mở bài đó. Đây là thay đổi có tác động lớn nhất ở backend."),
    ("B2. Thêm thời hạn chờ an toàn",
     "Nếu Google Drive phản hồi chậm, hệ thống sẽ hiển thị phần đã có thay vì treo toàn trang. Tránh trường hợp một tài liệu lỗi kéo sập cả trang."),
]
for t, d in B_items:
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xD9, 0x6B, 0x27)
    doc.add_paragraph(d, style="Body Text")

H("Mức C — Thay đổi kiến trúc (lâu dài)", 2)
P("Mục tiêu: loại bỏ tận gốc hiện tượng tải lại toàn trang mỗi lần chuyển bài, đưa hệ thống về đúng tiêu chuẩn web hiện đại.\n")

C_items = [
    ("C1. Chuyển sang điều hướng không tải lại (SPA-lite)",
     "Thay vì bấm nút là tải lại toàn trang, hệ thống sẽ chỉ thay phần nội dung bài học trong khi giữ nguyên khung trang. Giải quyết dứt điểm 5 việc nối tiếp đã nêu ở mục 2."),
]
for t, d in C_items:
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x7B, 0x2D, 0x18)
    doc.add_paragraph(d, style="Body Text")

# 5. Rủi ro nếu không xử lý
H("5. Nếu không xử lý thì sao?", 1)
BULLET([
    "Học viên tiếp tục gặp trải nghiệm kém, có thể ảnh hưởng đến uy tín khóa học và tỉ lệ duy trì.",
    "Khi số lượng bài học hoặc học viên tăng, tải máy chủ tăng theo cấp số nhân (do đang tải dư nhiều dữ liệu).",
    "Chi phí Google Drive API có thể tăng vì cùng một nội dung bị tải đi tải lại nhiều lần.",
    "Khi cần nâng cấp thêm tính năng (ví dụ ghi chú, đánh dấu bài), kiến trúc hiện tại sẽ khó mở rộng.",
])

# 6. Lộ trình
H("6. Lộ trình đề xuất (nếu được duyệt)", 1)
P("Đề xuất thứ tự triển khai, từng bước có thể đo hiệu quả trước khi đi tiếp:\n")
NUM([
    "Tuần 1: Thực hiện Mức A1 + A2 + A3 (độc lập nhau, không cần sửa backend). Đo lại cảm nhận.",
    "Tuần 2–3: Thực hiện Mức B1 — tách danh sách bài và công thức. Cần test cẩn thận vì đụng tới cả trang chủ khóa học.",
    "Tuần 4 trở đi: Lên kế hoạch cho Mức C1 — chuyển sang điều hướng không tải lại. Có thể đi kèm với một đợt nâng cấp giao diện lớn nếu đối tác muốn.",
])

CALLOUT(
    "Cam kết của đội kỹ thuật",
    "Không có bước nào chạm đến dữ liệu đang có (không migration, không xóa, không sửa cấu trúc database). Mỗi bước sẽ đo trước và sau để chứng minh tác động. Không triển khai bước nào nếu chưa có sự đồng ý của đối tác."
)

# 7. Phụ lục
H("Phụ lục: Cách đọc báo cáo kỹ thuật đầy đủ", 1)
P("Báo cáo chi tiết dành cho đội kỹ thuật nằm tại docs/LESSON_NAVIGATION_PERFORMANCE_INVESTIGATION.md trong cùng thư mục. Báo cáo đó có:", italic=True, size=10)
BULLET([
    "Sơ đồ luồng xử lý cụ thể (waterfall).",
    "Chỉ rõ từng dòng file và từng truy vấn cơ sở dữ liệu.",
    "Bảng ước lượng tác động và rủi ro cho từng đề xuất.",
    "Kế hoạch kiểm thử chi tiết trước khi triển khai.",
])
P("\nNếu đối tác muốn xem thêm bản kỹ thuật chi tiết, vui lòng liên hệ đội kỹ thuật.", italic=True, size=10)

# Save
doc.save(OUT)
print(f"WROTE: {OUT}")
print(f"Size: ", end="")
import os
print(f"{os.path.getsize(OUT):,} bytes")
